#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Rainer C. B. Herold

# Libraries
try:
    from ..VF import *
except ImportError:
    import sys
    sys.path.append('.')
    from VF import *

def Word_Table(Dict_Result):
    try:
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ModuleNotFoundError as e: Module_Error(f"The module was not found\n\n{e}\n\nPlease confirm with the button 'Return'")

    global Location
    document = Document()

    def Text_Format(Format, Size):
        font = document.styles[str(Format)].font
        font.name = 'TeleGrotesk Next'
        font.size = Pt(int(Size))

    def insertHR(paragraph): # Copied from Stackoverflow
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.insert_element_before(pBdr,
            'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
            'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
            'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
            'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
            'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
            'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
            'w:pPrChange'
        )
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), COLOR_Headline)
        pBdr.append(bottom)

    style = document.styles
    font_charstyle = style.add_style('TeleGrot', WD_STYLE_TYPE.CHARACTER)
    font = font_charstyle.font
    font.name = 'TeleGrotesk Next'
    font.size = Pt(25)
    font.color.rgb = RGBColor(65, 105, 225)

    Head_Result = document.add_paragraph('')
    Head_Result.add_run(f'Result - {Date}', style="TeleGrot")
    insertHR(Head_Result)
    Text_Format("Normal", 10)

    table = document.add_table(rows=len(Dict_Result)+1, cols=7)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text, hdr_cells[3].text, hdr_cells[4].text = 'URL', 'DNS', 'X-Frame-Options', 'X-XSS-Protection', 'Content-Security-Policy'
    hdr_cells[5].text, hdr_cells[6].text, hdr_cells[7].text = 'Strict-Transport-Security', 'X-Content-Type-Options', 'Referrer-Policy'

    n = 1
    for Target in Dict_Result['Header']:
        m = 0
        Cell = table.rows[n].cells
        Cell[m].text = Target
        for Result_Left, Result_Right in Dict_Result['Header'][Target].items():
            m += 1
            if ((Result_Left == "X-XSS-Protection" or Result_Left == Array_Header[1].lower() or Result_Left == Array_Header[1].isupper()) and (Result_Right == "1" or Result_Right == "1; mode=block") or Result_Right == "1; mode=BLOCK"): Result_Right = "FEHLT"
            elif ((Result_Left == "X-Content-Type-Options" or Result_Left == Array_Header[4].lower() or Result_Left == Array_Header[4].isupper()) and Result_Right != "nosniff" and Result_Right != "NOSNIFF"): Result_Right = "FEHLT"
            elif ((Result_Left == "X-Frame-Options" or Result_Left == Array_Header[0].lower() or Result_Left == Array_Header[0].isupper()) and Result_Right != "DENY" and Result_Right != "deny"): Result_Right = "FEHLT"

            if (Result_Right != "FEHLT"): Cell[m].text = "✓"
            else: Cell[m].text = "X"
        n += 1

    if (exists(join(Location, f'{File_Name}.docx'))):
        Question = input(f"The file already exists\n\n{e}\n\nDo you want to override it? (Y/N)")
        if (Question == 'Y' or Question == 'y'):
            Try_Remove_File(join(Location, f'{File_Name}.docx'))
            try: document.save(f'{File_Name}.docx')
            except OSError: document.save(join(getcwd(), f'{File_Name[:10]}{File_Name[19:]}.docx'))
        elif (Question == 'N' or Question == 'n'):
            n = 0
            for _, _, files in walk(Location, topdown=False):
                for file in files:
                    if (file.endswith('.docx') and 'result' in file): n += 1
            try: document.save(join(Location, f'{File_Name}_{n}.docx'))
            except OSError: document.save(join(getcwd(), f'{File_Name[:10]}{File_Name[19:]}_{n}.docx'))
        else: Error_Message('Your decision is not acceptable.')
    else:
        try: document.save(join(Location, f'{File_Name}.docx'))
        except OSError: document.save(join(getcwd(), f'{File_Name[:10]}{File_Name[19:]}.docx'))
